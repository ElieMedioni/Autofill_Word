import os
import pandas as pd
from docx import Document

# Fichier source Excel et modèle Word
EXCEL_FILE = "Data/format.xlsx"
TEMPLATE_WORD = "Data/template.docx"
OUTPUT_FOLDER = "Output"
EXCEL_RESULT = "Data/Textes_resultats.xlsx"

# Chargement des données
df = pd.read_excel(EXCEL_FILE)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

resultats = []

def replace_all_text(doc: Document, replacements: dict):
    # Paragraphes normaux
    for para in doc.paragraphs:
        if any(key in para.text for key in replacements):
            full_text = para.text
            for key, value in replacements.items():
                full_text = full_text.replace(key, str(value))
            for run in para.runs:
                run.text = ""
            para.runs[0].text = full_text
            

    """# Paragraphes dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if any(key in para.text for key in replacements):
                        full_text = para.text
                        for key, value in replacements.items():
                            full_text = full_text.replace(key, str(value))
                        for run in para.runs:
                            run.text = ""
                        para.runs[0].text = full_text"""



# Génération d'un document Word pour chaque ligne Excel
for idx, row in df.iterrows():
    try:
        name = str(row.iloc[0])
        number = str(row.iloc[2])

        safe_filename = "".join(c for c in f"{name}_{number}" if c.isalnum() or c in " _-").rstrip()

        print(f"Traitement de {safe_filename}")

        doc = Document(TEMPLATE_WORD)

        replacements = {f"<<COLONNE_{chr(65 + i)}>>": row.iloc[i] for i in range(len(row))}
        replace_all_text(doc, replacements)

        word_path = os.path.abspath(os.path.join(OUTPUT_FOLDER, f"{safe_filename}.docx"))
        doc.save(word_path)

        resultats.append("V")

    except Exception as e:
        print(f"Erreur pour la ligne {idx + 1} : {e}")
        resultats.append(f"Erreur : {str(e)}")

# Résumé des traitements
df["Document généré"] = resultats
df.to_excel(EXCEL_RESULT, index=False)

print(f"Tous les documents Word ont été générés dans '{OUTPUT_FOLDER}'. Résumé dans '{EXCEL_RESULT}'.")
